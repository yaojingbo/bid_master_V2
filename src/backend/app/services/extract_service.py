from __future__ import annotations
"""
Extract service for document element extraction using LLM.
"""
import asyncio
import contextlib
import io
import logging
import re
import time
from typing import AsyncGenerator, Dict, Any, Optional
import json

from app.services.llm_service import LLMService
from app.services.file_service import FileService
from app.infrastructure.pg_storage import add_extract, get_file, calculate_content_hash, find_completed_extract


def _format_elements_as_markdown(elements: list[dict[str, Any]], fallback_text: str = "") -> str:
    if elements:
        markdown = "\n\n".join(
            f"## {element.get('name', '要素')}\n{element.get('content', '')}".strip()
            for element in elements
            if isinstance(element, dict)
        ).strip()
        if markdown:
            return markdown
    return fallback_text.strip()


def _format_extract_result(record: dict) -> str:
    return _format_elements_as_markdown(record.get("elements") or [], str(record.get("content") or ""))

logger = logging.getLogger(__name__)
EXTRACTION_CACHE_VERSION = "pdf-table-ocr-docker-v1"


def _resolve_ocr_runtime_params(ocr_params: dict) -> tuple[str, list[int] | None, int | None, int]:
    engine = ocr_params.get("ocr_engine", "auto")
    page_numbers = ocr_params.get("ocr_pages")
    if "ocr_max_pages" in ocr_params:
        max_pages = ocr_params["ocr_max_pages"]
    elif engine == "ocrmypdf":
        max_pages = None
    elif engine == "llm":
        max_pages = 3
    else:
        max_pages = 15
    timeout_seconds = ocr_params.get("ocr_timeout_seconds") or (600 if engine == "ocrmypdf" else 120)
    return engine, page_numbers, max_pages, timeout_seconds


def _versioned_source_hash(source_hash: str, ocr_params: Optional[dict] = None) -> str:
    if not ocr_params:
        return f"{source_hash}:{EXTRACTION_CACHE_VERSION}"
    ocr_engine, ocr_pages, ocr_max_pages, _ = _resolve_ocr_runtime_params(ocr_params)
    pages_key = ",".join(str(page) for page in ocr_pages) if ocr_pages else ("all" if ocr_max_pages is None else f"first-{ocr_max_pages}")
    return f"{source_hash}:{EXTRACTION_CACHE_VERSION}:ocr={ocr_engine}:pages={pages_key}"


def _strip_json_fence(text: str) -> str:
    stripped = text.strip()
    if not stripped.startswith("```"):
        return stripped

    first_newline = stripped.find("\n")
    if first_newline == -1:
        return stripped

    body = stripped[first_newline + 1:]
    if body.rstrip().endswith("```"):
        body = body.rstrip()[:-3]
    return body.strip()


def _complete_json_candidate(text: str) -> str:
    stack: list[str] = []
    in_string = False
    escaped = False

    for char in text:
        if escaped:
            escaped = False
            continue
        if char == "\\" and in_string:
            escaped = True
            continue
        if char == '"':
            in_string = not in_string
            continue
        if in_string:
            continue
        if char in "[{":
            stack.append(char)
        elif char == "]" and stack and stack[-1] == "[":
            stack.pop()
        elif char == "}" and stack and stack[-1] == "{":
            stack.pop()

    suffix = "".join("}" if char == "{" else "]" for char in reversed(stack))
    return text + suffix


def _find_json_start(text: str) -> int:
    return min((idx for idx in (text.find("{"), text.find("[")) if idx != -1), default=-1)


def _find_balanced_json_end(text: str) -> int | None:
    stack: list[str] = []
    in_string = False
    escaped = False

    for index, char in enumerate(text):
        if escaped:
            escaped = False
            continue
        if char == "\\" and in_string:
            escaped = True
            continue
        if char == '"':
            in_string = not in_string
            continue
        if in_string:
            continue
        if char in "[{":
            stack.append(char)
        elif char == "]" and stack and stack[-1] == "[":
            stack.pop()
        elif char == "}" and stack and stack[-1] == "{":
            stack.pop()
        if not stack and char in "]}":
            return index + 1

    return None


def _coerce_json_value_to_content(value: Any) -> str:
    if isinstance(value, str):
        return value.strip()
    if isinstance(value, (dict, list)):
        return json.dumps(value, ensure_ascii=False, indent=2)
    if value is None:
        return ""
    return str(value).strip()


def _elements_from_json_mapping(parsed: dict[str, Any]) -> list[dict[str, str]]:
    elements = []
    for name, value in parsed.items():
        if name in {"summary", "metadata", "meta"}:
            continue
        content = _coerce_json_value_to_content(value)
        if content:
            elements.append({"name": str(name), "content": content})
    return elements


def _parse_llm_json_response(response: str) -> tuple[str, list[dict[str, Any]]]:
    candidate = _strip_json_fence(response)
    json_start = _find_json_start(candidate)
    if json_start == -1:
        raise json.JSONDecodeError("未找到 JSON 起始符", candidate, 0)

    candidate = candidate[json_start:].strip()
    json_end = _find_balanced_json_end(candidate)
    if json_end is not None:
        candidate = candidate[:json_end]

    try:
        parsed = json.loads(candidate)
        json_str = candidate
    except json.JSONDecodeError:
        candidate = _complete_json_candidate(candidate)
        parsed = json.loads(candidate)
        json_str = candidate

    if isinstance(parsed, dict) and "elements" in parsed:
        found_elements = parsed["elements"]
    elif isinstance(parsed, dict):
        mapped_elements = _elements_from_json_mapping(parsed)
        if mapped_elements:
            found_elements = mapped_elements
        else:
            found_elements = next(
                (v for v in parsed.values() if isinstance(v, list) and v and isinstance(v[0], dict)),
                [],
            )
    elif isinstance(parsed, list):
        found_elements = parsed
    else:
        found_elements = []

    return json_str, [e for e in found_elements if isinstance(e, dict)]


def _selected_element_names(elements: list[str] | None) -> list[str]:
    from app.services.prompt_builder import ELEMENT_NAMES
    return [ELEMENT_NAMES[element] for element in elements or [] if element in ELEMENT_NAMES]


def _split_by_selected_names(text: str, selected_names: list[str]) -> list[dict[str, str]]:
    if not selected_names:
        return []

    name_pattern = "|".join(re.escape(name) for name in sorted(selected_names, key=len, reverse=True))
    heading_pattern = re.compile(rf"^\s*(?:#{{1,6}}\s*)?(?:[-*]\s*)?(?:[一二三四五六七八九十\d]+[、.．:]\s*)?(?:\*\*)?({name_pattern})(?:\*\*)?\s*[:：]?\s*(.*)$", re.MULTILINE)
    matches = list(heading_pattern.finditer(text))
    if not matches:
        return []

    results = []
    for index, match in enumerate(matches):
        start = match.end()
        end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
        content = (match.group(2).strip() + "\n" + text[start:end].strip()).strip()
        if content:
            results.append({"name": match.group(1), "content": content})
    return results


def _parse_markdown_elements_response(response: str, selected_names: list[str] | None = None) -> list[dict[str, str]]:
    text = response.strip()
    if not text:
        return []

    selected_split = _split_by_selected_names(text, selected_names or [])
    if len(selected_split) > 1:
        return selected_split

    element_aliases = {
        "项目基本信息": ["项目基本信息", "基本信息", "项目概况"],
        "资质要求": ["资质要求", "资格要求", "资格条件", "资质门槛"],
        "业绩要求": ["业绩要求", "业绩门槛", "类似业绩", "企业业绩"],
        "人员要求": ["人员要求", "人员门槛", "项目负责人", "人员资格"],
        "评标办法": ["评标办法", "评审办法", "评分办法"],
        "分值分配与评分细则": ["分值分配与评分细则", "评分细则", "分值分配", "评分标准"],
        "定标方法": ["定标方法", "定标规则", "中标候选人确定"],
        "合同条款": ["合同条款", "合同主要条款", "付款方式", "工期要求"],
        "门槛要求": ["门槛要求", "投标门槛", "否决项", "资格门槛"],
    }
    alias_to_name = {
        alias: name
        for name, aliases in element_aliases.items()
        for alias in aliases
    }
    alias_pattern = "|".join(re.escape(alias) for alias in sorted(alias_to_name, key=len, reverse=True))
    heading_patterns = [
        re.compile(rf"^\s*(?:#{{1,6}}\s*)?(?:[-*]\s*)?(?:要素\s*)?[一二三四五六七八九十\d]+[.、：:]\s*(?:\*\*)?({alias_pattern})(?:\*\*)?\s*[:：]?\s*(.*)$"),
        re.compile(rf"^\s*(?:#{{1,6}}\s*)?(?:[-*]\s*)?(?:\*\*)?({alias_pattern})(?:\*\*)?\s*[:：]?\s*(.*)$"),
    ]

    grouped: dict[str, list[str]] = {name: [] for name in element_aliases}
    current_name: str | None = None
    current_lines: list[str] = []

    def flush_current() -> None:
        nonlocal current_name, current_lines
        if current_name and "\n".join(current_lines).strip():
            grouped[current_name].append("\n".join(current_lines).strip())
        current_name = None
        current_lines = []

    for line in text.splitlines():
        matched_name = None
        trailing_content = ""
        for pattern in heading_patterns:
            match = pattern.match(line)
            if match:
                matched_name = alias_to_name[match.group(1)]
                trailing_content = match.group(2).strip() if len(match.groups()) > 1 else ""
                break

        if matched_name:
            flush_current()
            current_name = matched_name
            if trailing_content:
                current_lines.append(trailing_content)
        elif current_name:
            current_lines.append(line)

    flush_current()

    elements = [
        {"name": name, "content": "\n\n---\n\n".join(parts)}
        for name, parts in grouped.items()
        if parts
    ]
    if elements:
        return elements

    sections = re.split(r"\n(?=#{1,6}\s+|[一二三四五六七八九十\d]+[、.．]\s+|\*\*[^*]{2,40}\*\*\s*$)", text)
    fallback_elements: list[dict[str, str]] = []
    for index, section in enumerate(sections, 1):
        section = section.strip()
        if not section:
            continue
        lines = section.splitlines()
        title = re.sub(r"^(#{1,6}\s*|[一二三四五六七八九十\d]+[、.．]\s*|\*\*)", "", lines[0]).strip("*：: ")
        content = "\n".join(lines[1:]).strip() or section
        if len(content) >= 20:
            fallback_elements.append({"name": title[:40] or f"分析结果{index}", "content": content})
    return fallback_elements[:8]


def _elements_from_plain_text(text: str, selected_names: list[str] | None = None) -> list[dict[str, str]]:
    stripped = text.strip()
    if not stripped:
        return []
    parsed = _parse_markdown_elements_response(stripped, selected_names)
    if parsed:
        return parsed
    if selected_names:
        return [{"name": selected_names[0], "content": stripped}]
    return [{"name": "分析结果", "content": stripped}]


def _normalize_elements(elements: list[dict[str, Any]], fallback_text: str, selected_names: list[str] | None = None) -> list[dict[str, Any]]:
    if not elements:
        return []

    normalized = []
    for element in elements:
        name = str(element.get("name") or "").strip()
        content = str(element.get("content") or "").strip()
        if selected_names and name not in selected_names:
            split_elements = _parse_markdown_elements_response(content or fallback_text, selected_names)
            if split_elements:
                normalized.extend(split_elements)
                continue
        normalized.append({"name": name or "分析结果", "content": content})

    if len(normalized) != 1:
        return normalized

    element = normalized[0]
    content = str(element.get("content") or fallback_text or "")
    split_elements = _parse_markdown_elements_response(content, selected_names)
    return split_elements if len(split_elements) > 1 else normalized


def _format_pdf_table(table: list[list[Any]]) -> str:
    rows = []
    for row in table:
        cells = [str(cell or "").strip() for cell in row]
        if any(cells):
            rows.append(" | ".join(cells))
    return "\n".join(rows)


def extract_text_from_pdf(content: bytes, max_pages: int = 120) -> tuple[str, bool]:
    """从PDF二进制内容中提取文本。返回 (文本, 是否需要OCR)。"""
    import pdfplumber
    parts = []
    page_count = 0
    low_text_pages: list[int] = []
    image_pages: list[int] = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for i, page in enumerate(pdf.pages[:max_pages], start=1):
            page_count += 1
            page_text = page.extract_text() or ""
            tables = page.extract_tables() or []
            table_parts = [_format_pdf_table(table) for table in tables]
            table_text = "\n\n".join(part for part in table_parts if part.strip())
            page_chars = len(page_text.strip()) + len(table_text.strip())
            has_images = bool(getattr(page, "images", None))

            if page_text.strip():
                parts.append(f"\n\n--- 第 {i} 页文本 ---\n{page_text.strip()}")
            if table_text.strip():
                parts.append(f"\n\n--- 第 {i} 页表格 ---\n{table_text.strip()}")
            if page_chars < 30:
                low_text_pages.append(i)
            if has_images:
                image_pages.append(i)

    text = "\n".join(parts)
    total_chars = len(text.strip())
    scanned_ratio = len(low_text_pages) / page_count if page_count else 0
    image_ratio = len(image_pages) / page_count if page_count else 0
    needs_ocr = page_count > 0 and (
        total_chars < 100
        or scanned_ratio >= 0.6
        or (image_ratio >= 0.5 and scanned_ratio >= 0.3)
    )
    logger.info(
        "PDF 文本提取: %d 页, %d 字符, 低文本页=%d, 图片页=%d, needs_ocr=%s",
        page_count,
        total_chars,
        len(low_text_pages),
        len(image_pages),
        needs_ocr,
    )
    return text, needs_ocr


def extract_text_from_docx(content: bytes) -> str:
    """从 .docx 二进制内容中提取文本（含段落和表格）"""
    from docx import Document
    doc = Document(io.BytesIO(content))
    parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(cells))

    return "\n".join(parts)


def extract_text_from_xlsx(content: bytes) -> str:
    """从 .xlsx 二进制内容中提取文本（按 sheet 组织）"""
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    parts = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"## Sheet: {sheet_name}")
        for row in ws.iter_rows(values_only=True):
            row_values = [str(c) if c is not None else "" for c in row]
            line = " | ".join(row_values).strip()
            if line:
                parts.append(line)
        parts.append("")

    wb.close()
    return "\n".join(parts)


def extract_text_from_xls(content: bytes) -> str:
    """从 .xls（旧格式）二进制内容中提取文本"""
    import xlrd
    wb = xlrd.open_workbook(file_contents=content)
    parts = []

    for sheet_name in wb.sheet_names():
        ws = wb.sheet_by_name(sheet_name)
        parts.append(f"## Sheet: {sheet_name}")
        for row_idx in range(ws.nrows):
            row_values = [str(ws.cell_value(row_idx, col_idx)) for col_idx in range(ws.ncols)]
            line = " | ".join(row_values).strip()
            if line:
                parts.append(line)
        parts.append("")

    return "\n".join(parts)


# OLE2 magic bytes (XLS, DOC)
_OLE2_MAGIC = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"


def extract_text_from_content(content: bytes) -> tuple[str, bool]:
    """根据文件内容自动判断格式并提取文本（magic bytes 检测）。返回 (文本, 是否需要OCR)。"""
    # PDF
    if content[:4] == b"%PDF":
        return extract_text_from_pdf(content)

    # ZIP-based: DOCX or XLSX
    if content[:4] == b"PK\x03\x04":
        import zipfile
        try:
            with zipfile.ZipFile(io.BytesIO(content)) as zf:
                names = zf.namelist()
                if any("word/" in n for n in names):
                    return extract_text_from_docx(content), False
                elif any("xl/" in n for n in names):
                    return extract_text_from_xlsx(content), False
        except Exception:
            pass
        return "", False

    # OLE2: XLS
    if content[:8] == _OLE2_MAGIC:
        try:
            return extract_text_from_xls(content), False
        except Exception:
            pass
        return "", False

    # Plain text (Markdown, CSV, TXT...)
    for encoding in ["utf-8", "gbk", "gb2312", "latin-1"]:
        try:
            return content.decode(encoding), False
        except (UnicodeDecodeError, ValueError):
            continue

    return "", False


class ExtractService:
    """Service for extracting elements from tender documents."""

    def __init__(
        self,
        llm_service: LLMService = None,
        file_service: FileService = None,
    ):
        self.llm = llm_service or LLMService()
        self.file_service = file_service or FileService()

    async def extract_elements_stream(
        self,
        document_id: str,
        provider: str = "deepseek",
        model: Optional[str] = None,
        template_type: str = "standard",
        elements: Optional[list[str]] = None,
        mode: str = "single",
        params: Optional[dict] = None,
        user_id: Optional[str] = None,
        request: Any = None,
    ) -> AsyncGenerator[Dict[str, Any], None]:
        """
        Extract elements from document with SSE streaming.

        Args:
            document_id: 文件 UUID
            provider: LLM 供应商
            model: 模型名
            template_type: standard | brief | batch | threshold
            elements: 指定提取的要素 key 列表
            mode: single | batch | threshold
            params: 额外参数（如 threshold 模式的 user_qualifications）
        """
        from app.services.prompt_builder import get_prompt_builder

        try:
            yield {"type": "progress", "message": "正在读取文档...", "phase": "reading", "percentage": 5}

            ocr_params = params or {}
            content = await self.file_service.download(document_id, user_id)
            source_hash = _versioned_source_hash(calculate_content_hash(content), ocr_params)
            cached = await find_completed_extract(source_hash, template_type, mode, elements, user_id) if user_id else None
            if cached:
                yield {"type": "progress", "message": "已复用同一原始文件的历史要素提取结果", "phase": "extracting", "percentage": 90}
                cached_elements = cached.get("elements") or []
                cached_text = _format_extract_result(cached)
                if cached_elements:
                    for element in cached_elements:
                        yield {"type": "element", "data": element, "phase": "extracting", "percentage": 90}
                elif cached_text:
                    yield {"type": "element", "data": {"name": "分析结果", "content": cached_text}, "phase": "extracting", "percentage": 90}
                yield {"type": "done", "data": {"summary": "已复用历史提取结果"}, "phase": "extracting", "percentage": 100}
                return

            text_content, needs_ocr = extract_text_from_content(content)

            # 图片型 PDF：触发 OCR 识别
            if needs_ocr:
                ocr_engine = (params or {}).get("ocr_engine", "auto")
                logger.info("触发 OCR: document_id=%s, engine=%s, provider=%s, model=%s", document_id, ocr_engine, provider, model)
                engine_hint = "本地 OCRmyPDF" if ocr_engine in ("auto", "ocrmypdf", "local") else ("本地 PaddleOCR" if ocr_engine == "paddleocr" else f"云端视觉模型 {provider}")
                yield {"type": "progress", "message": f"检测到图片型 PDF，正在启动 {engine_hint} 识别...", "phase": "parsing", "percentage": 12}
                try:
                    from app.services.ocr_service import ocr_pdf, resolve_vision_model
                    eff_provider, eff_model = resolve_vision_model(provider, model)
                    if (eff_provider, eff_model) != (provider, model):
                        logger.info("OCR 模型回退: %s/%s → %s/%s", provider, model, eff_provider, eff_model)
                        yield {"type": "progress", "message": f"当前模型不支持图片识别，已切换到 {eff_provider}/{eff_model}", "phase": "parsing", "percentage": 13}

                    ocr_result = None
                    async for ocr_event in self._run_ocr_with_progress(
                        content,
                        eff_provider,
                        eff_model,
                        user_id,
                        ocr_params,
                        request,
                    ):
                        if ocr_event.get("type") == "ocr_result":
                            ocr_result = ocr_event
                        else:
                            yield ocr_event

                    ocr_text = (ocr_result or {}).get("text", "")
                    logger.info("OCR 完成: %d 字符 (engine=%s)", len(ocr_text), ocr_engine)
                    if ocr_text.strip():
                        text_content = "\n\n".join(part for part in [text_content.strip(), ocr_text.strip()] if part)
                        yield {"type": "progress", "message": f"OCR 识别完成（{len(ocr_text)}字符），正在分析...", "phase": "parsing", "percentage": 20}
                    else:
                        yield {"type": "progress", "message": "OCR 未识别到文字，使用原始提取结果", "phase": "parsing", "percentage": 20}
                except Exception as e:
                    yield {"type": "progress", "message": f"OCR 识别失败（{str(e)[:80]}），使用原始提取结果", "phase": "parsing", "percentage": 20}

            if not text_content.strip():
                message = "文档内容为空或无法解析"
                if needs_ocr:
                    ocr_engine = (params or {}).get("ocr_engine", "auto")
                    message = f"扫描版 PDF 需要 OCR，但 {ocr_engine} 未能识别到文字；请检查 /api/health/ocr 或后端 OCR 环境"
                yield {"type": "error", "data": {"message": message}}
                return

            yield {"type": "progress", "message": f"文档解析完成（{len(text_content)}字符），正在分析...", "phase": "parsing", "percentage": 20}

            # 文档截断（200K 字符，约 50K token，与旧项目对齐）
            MAX_CHARS = 200000
            truncated = text_content[:MAX_CHARS]
            if len(text_content) > MAX_CHARS:
                truncated += f"\n\n[文档已截断，共 {len(text_content)} 字符]"

            # 用 PromptBuilder 构建 prompt（注入领域知识 + 提取规则 + 输出模板）
            builder = get_prompt_builder()
            system_prompt = builder.build_extract_system_prompt(template_type, elements)
            user_prompt = builder.build_extract_user_prompt(truncated, mode, params)

            messages = [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ]

            # Notify frontend that AI analysis is starting
            resolved_model = model or self.llm.llm.MODEL_MAP.get(provider, "unknown")
            if "/" in resolved_model:
                resolved_model = resolved_model.split("/", 1)[1]
            yield {"type": "progress", "message": f"AI 正在分析（{provider}/{resolved_model}）...", "phase": "analyzing", "percentage": 30}

            # LLM 流式调用 + 进度推送 + JSON 解析
            async for event in self._stream_llm_with_progress(
                provider, messages, model, document_id, template_type, user_id, source_hash=source_hash, result_mode=mode, elements=elements
            ):
                yield event

        except Exception as e:
            yield {"type": "error", "data": {"message": str(e)}}

    async def _run_ocr_with_progress(
        self,
        content: bytes,
        provider: str,
        model: Optional[str],
        user_id: Optional[str],
        ocr_params: dict,
        request: Any = None,
    ) -> AsyncGenerator[Dict[str, Any], None]:
        from app.services.ocr_service import ocr_pdf

        queue: asyncio.Queue = asyncio.Queue()
        cancel_event = asyncio.Event()
        started_at = time.monotonic()
        result_holder = {"text": "", "error": None}
        engine = ocr_params.get("ocr_engine", "auto")
        page_numbers = ocr_params.get("ocr_pages")
        if "ocr_max_pages" in ocr_params:
            max_pages = ocr_params["ocr_max_pages"]
        elif engine == "ocrmypdf":
            max_pages = None
        else:
            max_pages = 15
        timeout_seconds = ocr_params.get("ocr_timeout_seconds") or (600 if engine == "ocrmypdf" else 120)
        last_heartbeat = started_at

        async def progress_callback(page: int, total: int):
            percentage = min(25, 12 + int((page / max(total, 1)) * 13))
            await queue.put({
                "type": "ocr_progress",
                "message": f"OCR 正在识别第 {page}/{total} 页...",
                "phase": "parsing",
                "percentage": percentage,
                "currentPage": page,
                "totalPages": total,
            })

        async def background_task():
            try:
                result_holder["text"] = await ocr_pdf(
                    content,
                    provider,
                    model,
                    user_id,
                    progress_callback=progress_callback,
                    page_numbers=page_numbers,
                    engine=engine,
                    max_pages=max_pages,
                    timeout_seconds=timeout_seconds,
                    cancel_event=cancel_event,
                )
                await queue.put({"type": "ocr_done"})
            except asyncio.CancelledError:
                raise
            except Exception as e:
                result_holder["error"] = str(e)
                await queue.put({"type": "ocr_error", "error": str(e)})

        task = asyncio.create_task(background_task())
        await queue.put({
            "type": "ocr_progress",
            "message": "OCR 任务已启动，正在准备识别引擎...",
            "phase": "parsing",
            "percentage": 13,
            "currentPage": 0,
            "totalPages": max_pages,
        })
        try:
            while True:
                if time.monotonic() - started_at > timeout_seconds:
                    cancel_event.set()
                    task.cancel()
                    yield {"type": "progress", "message": f"OCR 超过 {timeout_seconds} 秒未完成，已停止并使用原始解析结果", "phase": "parsing", "percentage": 20}
                    return

                if request:
                    disconnected = False
                    with contextlib.suppress(Exception):
                        disconnected = await asyncio.wait_for(request.is_disconnected(), timeout=0.05)
                    if disconnected:
                        cancel_event.set()
                        task.cancel()
                        yield {"type": "progress", "message": "OCR 已取消", "phase": "parsing", "percentage": 20}
                        return

                try:
                    event = await asyncio.wait_for(queue.get(), timeout=1.0)
                except asyncio.TimeoutError:
                    if time.monotonic() - last_heartbeat > 2.0:
                        last_heartbeat = time.monotonic()
                        yield {"type": "ocr_progress", "message": "OCR 正在处理中，请稍候...", "phase": "parsing", "percentage": 14}
                    continue

                if event["type"] == "ocr_done":
                    with contextlib.suppress(Exception):
                        from app.services.ocrmypdf_service import get_last_ocrmypdf_evidence

                        evidence = get_last_ocrmypdf_evidence()
                        if evidence:
                            yield {
                                "type": "ocr_progress",
                                "message": f"本地 OCR 证据已保存：PDF {evidence['pdf_path']}；文本 {evidence['text_path']}",
                                "phase": "parsing",
                                "percentage": 19,
                                "evidence": evidence,
                            }
                    yield {"type": "ocr_result", "text": result_holder["text"]}
                    return
                if event["type"] == "ocr_error":
                    raise RuntimeError(event["error"])
                yield event
        finally:
            if not task.done():
                cancel_event.set()
                task.cancel()

    async def _stream_llm_with_progress(
        self,
        provider: str,
        messages: list[dict],
        model: Optional[str],
        file_id: str,
        template_type: str,
        user_id: Optional[str] = None,
        source_hash: Optional[str] = None,
        result_mode: Optional[str] = None,
        elements: Optional[list[str]] = None,
    ) -> AsyncGenerator[Dict[str, Any], None]:
        """LLM 流式调用 + 进度推送 + JSON 解析。"""
        import asyncio
        chunk_queue: asyncio.Queue = asyncio.Queue()
        result_holder = {"text": "", "done": False, "error": None}
        _saved = False

        file_record = await get_file(file_id, user_id)
        file_name = file_record.get("original_name", "") if file_record else ""
        effective_source_hash = source_hash or (file_record.get("file_hash") if file_record else None)
        effective_mode = result_mode or provider
        selected_names = _selected_element_names(elements)

        resolved_model = model or self.llm.llm.MODEL_MAP.get(provider, "unknown")
        if "/" in resolved_model:
            resolved_model = resolved_model.split("/", 1)[1]

        async def _llm_background_task():
            """后台运行 LLM 调用，chunks 推入 Queue。即使 SSE 断连也会继续运行。"""
            try:
                count = 0
                async for chunk in self.llm.llm.complete(provider, messages, model=model, stream=True, user_id=user_id):
                    result_holder["text"] += chunk
                    count += 1
                    await chunk_queue.put({"type": "chunk", "count": count})
                result_holder["done"] = True
                await chunk_queue.put({"type": "llm_done"})
            except Exception as e:
                result_holder["error"] = str(e)
                await chunk_queue.put({"type": "llm_error", "error": str(e)})

        background_task = asyncio.create_task(_llm_background_task())

        try:
            last_pct = 30
            while True:
                try:
                    event = await asyncio.wait_for(chunk_queue.get(), timeout=2.0)
                except asyncio.TimeoutError:
                    if not result_holder["done"] and not result_holder["error"]:
                        last_pct = min(last_pct + 1, 88)
                        yield {"type": "llm_progress", "message": f"AI 正在分析（{provider}/{resolved_model}）... {last_pct}%", "phase": "analyzing", "percentage": last_pct}
                    continue

                if event["type"] == "chunk":
                    if event["count"] % 5 == 0:
                        pct = min(30 + event["count"] // 2, 88)
                        last_pct = max(last_pct, pct)
                        yield {"type": "llm_progress", "message": f"AI 正在分析（{provider}/{resolved_model}）... {last_pct}%", "phase": "analyzing", "percentage": last_pct}
                elif event["type"] == "llm_done":
                    # LLM 完成，解析并保存
                    full_response = result_holder["text"]
                    try:
                        json_str, found_elements = _parse_llm_json_response(full_response)
                        found_elements = _normalize_elements(found_elements, full_response, selected_names)

                        markdown_content = _format_elements_as_markdown(found_elements, full_response)
                        await add_extract({
                            "file_id": file_id,
                            "file_name": file_name,
                            "template_type": template_type,
                            "mode": effective_mode,
                            "content": markdown_content,
                            "elements": found_elements,
                            "status": "completed",
                            "source_hash": effective_source_hash,
                        }, user_id=user_id)
                        _saved = True

                        for element in found_elements:
                            yield {"type": "element", "data": element, "phase": "extracting", "percentage": 90}
                        if not found_elements and full_response.strip():
                            yield {"type": "element", "data": {"name": "分析结果", "content": full_response.strip()}, "phase": "extracting", "percentage": 90}
                        yield {
                            "type": "done",
                            "data": {"summary": "文档分析完成", "elementCount": max(len(found_elements), 1 if full_response.strip() else 0)},
                            "phase": "extracting",
                            "percentage": 100,
                        }
                    except json.JSONDecodeError:
                        found_elements = _elements_from_plain_text(full_response, selected_names)
                        content = _format_elements_as_markdown(found_elements, full_response)
                        await add_extract({
                            "file_id": file_id,
                            "file_name": file_name,
                            "template_type": template_type,
                            "mode": effective_mode,
                            "content": content,
                            "elements": found_elements,
                            "status": "completed",
                            "source_hash": effective_source_hash,
                        }, user_id=user_id)
                        _saved = True
                        for element in found_elements:
                            yield {"type": "element", "data": element, "phase": "extracting", "percentage": 90}
                        yield {
                            "type": "done",
                            "data": {"summary": "文档分析完成", "elementCount": len(found_elements)},
                            "phase": "extracting",
                            "percentage": 100,
                        }
                    break
                elif event["type"] == "llm_error":
                    yield {"type": "error", "data": {"message": event["error"]}}
                    break

        finally:
            if not _saved:
                if not background_task.done():
                    background_task.cancel()
                    import contextlib
                    with contextlib.suppress(asyncio.CancelledError):
                        await background_task

                if result_holder["done"]:
                    full_response = result_holder["text"]
                    await add_extract({
                        "file_id": file_id,
                        "file_name": file_name,
                        "template_type": template_type,
                        "mode": effective_mode,
                        "content": full_response,
                        "status": "completed_disconnected",
                        "source_hash": effective_source_hash,
                    }, user_id=user_id)
                elif result_holder["text"]:
                    await add_extract({
                        "file_id": file_id,
                        "file_name": file_name,
                        "template_type": template_type,
                        "mode": effective_mode,
                        "content": result_holder["text"],
                        "status": "partial",
                        "source_hash": effective_source_hash,
                    }, user_id=user_id)

    async def extract_batch_stream(
        self,
        file_ids: list[str],
        provider: str = "deepseek",
        model: Optional[str] = None,
        elements: Optional[list[str]] = None,
        user_id: Optional[str] = None,
        ocr_params: Optional[dict] = None,
    ) -> AsyncGenerator[Dict[str, Any], None]:
        """批量对比：并行下载多个文件 → 合并 → batch 模板 → LLM 对比。"""
        from app.services.prompt_builder import get_prompt_builder
        import asyncio

        try:
            ocr_params = ocr_params or {}
            ocr_engine = ocr_params.get("ocr_engine", "auto")
            total = len(file_ids)
            yield {"type": "progress", "message": f"正在读取 {total} 个文件...", "phase": "reading", "percentage": 5}

            # 并行下载 + 解析所有文件
            async def fetch_and_parse(fid: str) -> tuple[str, str]:
                try:
                    content = await self.file_service.download(fid, user_id)
                    text, needs_ocr = extract_text_from_content(content)
                    if needs_ocr:
                        try:
                            from app.services.ocr_service import ocr_pdf, resolve_vision_model
                            eff_provider, eff_model = resolve_vision_model(provider, model)
                            ocr_text = await ocr_pdf(
                                content,
                                eff_provider,
                                eff_model,
                                user_id,
                                page_numbers=ocr_params.get("ocr_pages"),
                                engine=ocr_engine,
                                max_pages=_resolve_ocr_runtime_params(ocr_params)[2],
                                timeout_seconds=_resolve_ocr_runtime_params(ocr_params)[3],
                            )
                            if ocr_text.strip():
                                text = "\n\n".join(part for part in [text.strip(), ocr_text.strip()] if part)
                        except Exception as ocr_err:
                            logger.warning("批量提取 OCR 失败: file_id=%s, error=%s", fid, ocr_err)
                    return fid, text
                except Exception as e:
                    return fid, f"[读取失败: {e}]"

            results = await asyncio.gather(*[fetch_and_parse(fid) for fid in file_ids])

            # 合并文本
            all_parts = []
            success_count = 0
            for fid, text in results:
                if not text.strip():
                    all_parts.append(f"## 文件: {fid}\n[内容为空]\n")
                elif text.startswith("[读取失败"):
                    all_parts.append(f"## 文件: {fid}\n{text}\n")
                else:
                    success_count += 1
                    all_parts.append(f"## 文件: {fid}\n{text[:50000]}\n")

            if success_count < 2:
                yield {"type": "error", "data": {"message": f"至少需要 2 个有效文件，当前仅有 {success_count} 个"}}
                return

            combined = "\n\n---\n\n".join(all_parts)
            MAX_CHARS = 200000
            if len(combined) > MAX_CHARS:
                combined = combined[:MAX_CHARS] + f"\n\n[内容已截断，共 {len(combined)} 字符]"

            yield {"type": "progress", "message": f"{success_count} 个文件读取完成，正在对比分析...", "phase": "parsing", "percentage": 20}

            builder = get_prompt_builder()
            system_prompt = builder.build_extract_system_prompt("batch", elements)
            user_prompt = builder.build_extract_user_prompt(combined, mode="batch")

            messages = [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ]

            yield {"type": "progress", "message": "AI 正在对比分析多份招标文件...", "phase": "analyzing", "percentage": 30}

            async for event in self._stream_llm_with_progress(
                provider, messages, model, ",".join(file_ids), "batch", user_id, elements=elements
            ):
                yield event

        except Exception as e:
            yield {"type": "error", "data": {"message": str(e)}}

    async def extract_threshold_stream(
        self,
        file_id: str,
        user_qualifications: str,
        provider: str = "deepseek",
        model: Optional[str] = None,
        user_id: Optional[str] = None,
        ocr_params: Optional[dict] = None,
    ) -> AsyncGenerator[Dict[str, Any], None]:
        """门槛分析：下载文件 → threshold 模板 + 用户资质 → LLM 逐项比对。"""
        from app.services.prompt_builder import get_prompt_builder

        try:
            ocr_params = ocr_params or {}
            ocr_engine = ocr_params.get("ocr_engine", "auto")
            yield {"type": "progress", "message": "正在读取招标文件...", "phase": "reading", "percentage": 5}

            content = await self.file_service.download(file_id, user_id)
            source_hash = _versioned_source_hash(calculate_content_hash(content), ocr_params)
            cached = await find_completed_extract(source_hash, "standard", "single", None, user_id) if user_id else None
            if cached:
                cached_text = _format_extract_result(cached)
                yield {"type": "progress", "message": "已复用历史要素提取结果，正在比对门槛要求...", "phase": "parsing", "percentage": 25}
                text_content = cached_text
                needs_ocr = False
            else:
                text_content, needs_ocr = extract_text_from_content(content)

            if needs_ocr:
                logger.info("门槛分析触发 OCR: file_id=%s, engine=%s, provider=%s", file_id, ocr_engine, provider)
                engine_hint = "本地 OCRmyPDF" if ocr_engine in ("auto", "ocrmypdf", "local") else ("本地 PaddleOCR" if ocr_engine == "paddleocr" else f"云端视觉模型 {provider}")
                yield {"type": "progress", "message": f"检测到图片型 PDF，正在启动 {engine_hint} 识别...", "phase": "parsing", "percentage": 12}
                try:
                    from app.services.ocr_service import ocr_pdf, resolve_vision_model
                    eff_provider, eff_model = resolve_vision_model(provider, model)
                    if (eff_provider, eff_model) != (provider, model):
                        logger.info("OCR 模型回退: %s/%s → %s/%s", provider, model, eff_provider, eff_model)
                        yield {"type": "progress", "message": f"当前模型不支持图片识别，已切换到 {eff_provider}/{eff_model}", "phase": "parsing", "percentage": 13}
                    ocr_engine, ocr_pages, ocr_max_pages, ocr_timeout_seconds = _resolve_ocr_runtime_params(ocr_params)
                    ocr_text = await ocr_pdf(
                        content,
                        eff_provider,
                        eff_model,
                        user_id,
                        page_numbers=ocr_pages,
                        engine=ocr_engine,
                        max_pages=ocr_max_pages,
                        timeout_seconds=ocr_timeout_seconds,
                    )
                    logger.info("门槛 OCR 完成: %d 字符 (engine=%s)", len(ocr_text), ocr_engine)
                    if ocr_text.strip():
                        text_content = ocr_text
                except Exception:
                    pass  # OCR 失败不影响主流程

            if not text_content.strip():
                message = "文档内容为空或无法解析"
                if needs_ocr:
                    message = f"扫描版 PDF 需要 OCR，但 {ocr_engine} 未能识别到文字；请检查 /api/health/ocr 或后端 OCR 环境"
                yield {"type": "error", "data": {"message": message}}
                return

            yield {"type": "progress", "message": f"文档解析完成（{len(text_content)}字符），正在比对分析...", "phase": "parsing", "percentage": 20}

            MAX_CHARS = 200000
            truncated = text_content[:MAX_CHARS]

            builder = get_prompt_builder()
            system_prompt = builder.build_extract_system_prompt("threshold")
            user_prompt = builder.build_extract_user_prompt(
                truncated, mode="threshold", params={"user_qualifications": user_qualifications}
            )

            messages = [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ]

            yield {"type": "progress", "message": "AI 正在逐项比对门槛要求...", "phase": "analyzing", "percentage": 30}

            async for event in self._stream_llm_with_progress(
                provider, messages, model, file_id, "threshold", user_id, elements=["qualification", "experience", "personnel"]
            ):
                yield event

        except Exception as e:
            yield {"type": "error", "data": {"message": str(e)}}


_extract_service: ExtractService | None = None


def get_extract_service() -> ExtractService:
    """Get or create extract service instance."""
    global _extract_service
    if _extract_service is None:
        _extract_service = ExtractService()
    return _extract_service